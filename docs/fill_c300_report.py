"""Fill AY2026 C300 Report Template with JobPortal content."""
from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

SRC = Path(r"c:\Users\rioan\Downloads\AY2026 C300 Report Template.docx")
OUT1 = Path(r"c:\Users\rioan\Downloads\AY2026_C300_JobPortal_Report_DRAFT.docx")
OUT2 = Path(r"c:\Y3S1\FYP\jobportal\docs\AY2026_C300_JobPortal_Report_DRAFT.docx")


def set_runs(p, text, bold=None):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
        if bold is not None:
            p.runs[0].bold = bold
    else:
        run = p.add_run(text)
        if bold is not None:
            run.bold = bold


def clear_para(p):
    set_runs(p, "")


def insert_after(paragraph, text):
    new_p = deepcopy(paragraph._p)
    for child in list(new_p):
        if child.tag == qn("w:r"):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    p = Paragraph(new_p, paragraph._parent)
    set_runs(p, text)
    return p


def find_para(doc, startswith=None, equals=None, contains=None):
    for p in doc.paragraphs:
        t = p.text.strip()
        if equals is not None and t == equals:
            return p
        if startswith is not None and t.startswith(startswith):
            return p
        if contains is not None and contains in t:
            return p
    return None


def main():
    shutil.copy2(SRC, OUT1)
    doc = Document(str(OUT1))

    # Cover
    set_runs(doc.paragraphs[3], "JobPortal — Full-Stack Job Portal Web Application")
    set_runs(doc.paragraphs[4], "<Project ID — fill in from Supervisor / module coordinator>")
    set_runs(doc.paragraphs[7], "Date of Submission: 27-Jul-2026")
    set_runs(doc.paragraphs[13], "<Team ID — fill in>")

    # Acknowledgements
    set_runs(
        doc.paragraphs[17],
        "We would like to thank our Final Year Project supervisor for continuous guidance, "
        "constructive feedback during project meetings, and support in refining the project scope. "
        "We also thank the lecturers and staff of the School of Infocomm for providing the learning "
        "foundation in web development, databases, and software engineering practices that enabled "
        "this project. Finally, we appreciate classmates and peers who tested early prototypes of "
        "JobPortal and shared usability feedback that improved the student, employer, and admin experiences.",
    )

    # Abstract
    set_runs(
        doc.paragraphs[47],
        "JobPortal is a full-stack web application that connects students seeking employment opportunities "
        "with employers posting jobs, under administrator oversight. The platform addresses fragmented job "
        "discovery for students and inefficient applicant management for employers by providing a single, "
        "role-based portal for browsing, applying, posting, moderating, and analysing job activity. "
        "The system is implemented using Node.js and Express for a REST API, MySQL (Aiven) for persistence, "
        "and a Bootstrap 5 with vanilla JavaScript frontend. Security uses JWT authentication, bcrypt password "
        "hashing, and role-based access control for student, employer, and admin roles.",
    )
    set_runs(
        doc.paragraphs[48],
        "Key capabilities include job listing with search and multi-filter, bookmarks, applications with cover "
        "letter and resume (PDF upload or URL), employer job posting and applicant pipeline management, admin "
        "moderation and Chart.js analytics, in-app notifications for application status changes and interest-based "
        "job alerts, employer verification, job approval workflow, expired-job auto-close, and optional Gemini AI "
        "chat/smart search via a secured server endpoint.",
    )
    set_runs(
        doc.paragraphs[49],
        "The project accomplished a complete demo-ready portal with sample data, supporting design documentation "
        "(requirements, architecture, ERD, use cases, workflow, Gantt chart), and enhanced employer/admin dashboards. "
        "Final status is functionally complete for FYP demonstration. Recommended next steps are email delivery for "
        "notifications, password reset, public company profiles, broader automated testing, and production deployment.",
    )
    clear_para(doc.paragraphs[50])
    clear_para(doc.paragraphs[51])

    # Introduction
    set_runs(
        doc.paragraphs[53],
        "Many students struggle to find relevant internships and entry-level roles because opportunities are "
        "scattered across multiple platforms, while employers lack a simple way to post campus-oriented jobs "
        "and track applicants in one place. JobPortal solves this by offering a dedicated web portal tailored "
        "to three stakeholders: students (job seekers), employers (job posters), and administrators (moderators).",
    )
    set_runs(
        doc.paragraphs[54],
        "Motivation: Improving early-career job matching matters for student employability and for employers "
        "who need efficient, verified hiring pipelines. A moderated portal also reduces spam and low-quality "
        "listings through admin approval and employer verification.",
    )
    set_runs(
        doc.paragraphs[55],
        "Scope / major deliverables: (1) authenticated multi-role web application; (2) job browse/search/apply "
        "and bookmark flows; (3) employer posting and applicant status management; (4) admin dashboard with "
        "user/job management and analytics; (5) supporting design artefacts and user/technical documentation; "
        "(6) enhancements such as notifications, resume upload, job alerts, and AI-assisted search/chat.",
    )
    set_runs(
        doc.paragraphs[56],
        "Approach: Requirements were gathered and documented by role/module; the data model was designed in MySQL; "
        "a layered architecture (browser → Express API → MySQL) was implemented iteratively; features were tested "
        "manually against functional requirements; documentation pages and this report capture design and usage.",
    )
    set_runs(
        doc.paragraphs[57],
        "Assumptions: Users have a modern browser and internet access; MySQL credentials and JWT secret are "
        "configured via environment variables; employers provide accurate company/job information; admins actively "
        "moderate pending jobs and employer verification; Gemini AI features require a valid API key and are optional.",
    )
    set_runs(
        doc.paragraphs[58],
        "Main results: A working JobPortal at localhost:5000 with student/employer/admin journeys, REST APIs, "
        "cloud MySQL schema (users, employers, jobs, applications, bookmarks, job_views, notifications), and "
        "demo accounts for evaluation.",
    )
    clear_para(doc.paragraphs[59])

    # Section 2 intro
    set_runs(
        doc.paragraphs[62],
        "Agreed deliverables with Supervisor (summary): Business analysis covering student/employer pain points "
        "and a process-flow solution; full-stack technical design and implementation of JobPortal; analysis tools "
        "including use cases, ERD, architecture diagrams, and Gantt planning; resources comprising Node.js/Express "
        "server, Bootstrap frontend, and Aiven MySQL. Optional UML class/sequence diagrams for OOP were omitted "
        "as this project uses a procedural/REST style rather than a heavy OOP domain model.",
    )
    for i in [63, 64, 65, 66, 67, 69, 70, 71, 72]:
        clear_para(doc.paragraphs[i])

    # 2.1 Overview
    set_runs(
        doc.paragraphs[75],
        "JobPortal is a web-based job marketplace for students and employers. Students register, build profiles "
        "(skills, bio, interest fields, resume), browse and filter jobs, bookmark listings, and apply with a cover "
        "letter. Employers register company profiles (including UEN and verification), post jobs pending admin "
        "approval, manage applicants, and view dashboard insights. Administrators moderate users and jobs, approve "
        "or reject listings with reasons, verify employers, and monitor platform analytics.",
    )
    set_runs(
        doc.paragraphs[76],
        "Project Motivation: Centralise student-focused job discovery; give employers a structured posting and "
        "applicant workflow; give admins visibility and control to keep the marketplace trustworthy. Benefits "
        "include faster job discovery, clearer application tracking, reduced spam via moderation, and data-driven "
        "insights (views, applications, conversion).",
    )
    set_runs(
        doc.paragraphs[77],
        "Project Objectives: (1) Deliver secure role-based authentication; (2) enable end-to-end job apply/post/manage "
        "flows; (3) provide admin analytics and moderation; (4) improve UX with notifications, alerts, and optional AI "
        "assistance; (5) document architecture, requirements, and installation for maintainability.",
    )
    set_runs(
        doc.paragraphs[78],
        "Project Scope: In scope — web UI + REST API + MySQL data layer for student/employer/admin features listed "
        "above, design docs, testing evidence, and report. Out of scope — native mobile apps, payment/payroll, "
        "full ATS integrations, and mass email marketing automation (in-app notifications are included instead).",
    )
    set_runs(
        doc.paragraphs[79],
        "Project Assumptions: Single deployment instance for FYP demo; English UI; Singapore-oriented location terms "
        "(e.g. Remote, Hybrid, Islandwide); seed/demo data acceptable for evaluation; Gen AI chat available only when "
        "GEMINI_API_KEY is configured.",
    )

    # 2.2 Functional requirements
    set_runs(
        doc.paragraphs[82],
        "Functional requirements are organised by role (aligned with the project Project Specification page).\n\n"
        "Student / Job Seeker: Register and login; browse paginated jobs; search and filter; view job details; apply "
        "with cover letter and resume (upload PDF or URL); bookmark/unbookmark; track application status; edit profile "
        "(bio, skills, interest fields); receive notifications for status changes and matching job alerts.\n\n"
        "Employer: Register company; complete company profile; post jobs (type, category, salary, deadline, location); "
        "edit jobs; view applicants; update application status (pending/reviewed/shortlisted/rejected/hired); view "
        "dashboard insights (recent applicants, status breakdown, job performance, jobs closing soon); filter own jobs "
        "(approved/open/past/pending/rejected).\n\n"
        "Admin: Login; view dashboard stats; manage users (filter by role; soft-remove with reason); manage jobs "
        "(approve/reject/close with reasons); verify employers; view analytics charts (category/type, daily views, "
        "top applied jobs).\n\n"
        "System / Non-functional business rules: JWT required for protected routes; passwords hashed with bcrypt; "
        "jobs start as pending until admin approval; expired deadlines auto-close; removed accounts cannot log in "
        '("This account is invalid."); AI chat is rate-limited per user and globally; resume uploads restricted to PDF.',
    )

    # 2.3 Project plan
    set_runs(
        doc.paragraphs[84],
        "The project followed a 14-week plan (20 April 2026 – 27 July 2026), summarised from the team Gantt chart.\n\n"
        "Phase 1 — Planning & Requirements (W1–W2): scoping, requirements gathering, ERD/use cases, tech stack decision.\n"
        "Phase 2 — UI/UX Design (W2–W4): wireframes, schema design, Bootstrap frontend shells.\n"
        "Phase 3 — Development (W4–W8): Express/MySQL APIs, JWT auth, jobs/search, bookmarks/applications, employer "
        "posting, admin dashboard/analytics, then enhancements (notifications, resume upload, verification, AI search).\n"
        "Phase 4 — Testing & Bug Fixes (W9–W12): API/integration testing, UI polish, UAT.\n"
        "Phase 5 — Deployment & Submission (W8–W14): Aiven MySQL deployment, final documentation/report, final submission "
        "milestone on 27 July 2026.\n\n"
        "Work allocation: [FILL IN — assign each member fair ownership, e.g. Member A: backend auth/jobs APIs; "
        "Member B: frontend student/employer pages; Member C: admin dashboard & analytics; Member D: documentation, "
        "testing, deployment]. Ensure allocation matches actual Git commits and meeting minutes.\n\n"
        "Milestone alignment: Meeting 1 — proposal/overview; Meeting 2 — specification & plan (this section); "
        "Meeting 3 — design/implementation progress; Final — testing, docs, demo, report.",
    )
    for i in [86, 87, 88, 89]:
        clear_para(doc.paragraphs[i])

    # Business analysis
    set_runs(
        doc.paragraphs[92],
        "This section analyses the business problem JobPortal addresses for students, employers, and platform operators, "
        "and how the IT solution supports those processes.",
    )
    clear_para(doc.paragraphs[94])
    clear_para(doc.paragraphs[95])
    clear_para(doc.paragraphs[96])

    set_runs(
        doc.paragraphs[98],
        "Current situation: Students often jump between LinkedIn, company career pages, Telegram groups, and notice boards. "
        "Application status is opaque. Employers may post informally and track applicants in spreadsheets. Without "
        "moderation, portals risk fake employers and low-quality listings.",
    )
    set_runs(
        doc.paragraphs[99],
        "Business issues: (1) Fragmented job discovery increases student search time; (2) Weak application tracking "
        "reduces follow-through; (3) Employers lack a lightweight ATS-like view of applicants; (4) Platform trust requires "
        "verification and listing approval; (5) Expired jobs remaining open waste applicant effort.",
    )
    clear_para(doc.paragraphs[100])

    set_runs(
        doc.paragraphs[103],
        "Market segment: The primary segment is tertiary students seeking internships, part-time, and entry-level roles, "
        "plus SMEs and campus recruiters hiring locally (Singapore-focused locations and job types). Secondary segment is "
        "institution/admin operators who need oversight and analytics.",
    )
    set_runs(
        doc.paragraphs[104],
        "Competitive analysis (high level): General job boards (e.g. LinkedIn, Indeed, JobStreet) are broad but not "
        "campus-workflow focused; school career portals may be closed or limited. JobPortal differentiates with an "
        "integrated three-role design, admin moderation, interest-based alerts, and FYP-demonstrable analytics/AI assist, "
        "while remaining lightweight to deploy and explain.",
    )
    clear_para(doc.paragraphs[105])

    set_runs(
        doc.paragraphs[108],
        "Business solution: A single web portal with clear process flows.\n\n"
        "Student flow: Register → complete profile/resume/interests → search/filter jobs → bookmark or apply → "
        "receive status notifications → track outcomes on profile.\n\n"
        "Employer flow: Register company → await/complete verification → post job → admin approves → applicants apply → "
        "employer reviews and updates status (shortlist/reject/hire) → dashboard shows performance.\n\n"
        "Admin flow: Review pending employers/jobs → approve or reject with reason → monitor users/jobs/analytics → "
        "soft-remove abusive accounts.\n\n"
        "How IT helps: Centralised data (MySQL) keeps jobs/applications consistent; REST APIs enable responsive UI; "
        "JWT RBAC enforces least privilege; notifications close the feedback loop; analytics turn views/applications "
        "into actionable insights; auto-close of expired jobs keeps listings trustworthy.",
    )
    clear_para(doc.paragraphs[109])
    clear_para(doc.paragraphs[110])

    # Design
    set_runs(
        doc.paragraphs[112],
        "JobPortal uses a classic three-tier web architecture. Clients run in the browser; the application tier is a "
        "Node.js Express server exposing JSON APIs and serving static frontend files; the data tier is MySQL on Aiven. "
        "Optional external AI (Google Gemini) is called only from the server.",
    )
    set_runs(
        doc.paragraphs[114],
        "System architecture:\n"
        "1) Presentation: HTML/CSS/Bootstrap 5 pages (index, login/register, jobs, profile, post-job, employer, admin) "
        "and client JS (api.js, admin.js, chatbot.js).\n"
        "2) Application: Express routes — /api/auth, /api/jobs, /api/applications, /api/bookmarks, /api/admin, "
        "/api/notifications, /api/chat — with JWT middleware and multer for resume uploads under /uploads/resumes.\n"
        "3) Data: MySQL tables — users, employers, jobs, applications, bookmarks, job_views, notifications.\n"
        "4) External services: Aiven MySQL (TLS); Google Gemini API for chat/smart search (server-side key).\n\n"
        "Network topology for demo: User browser → HTTP localhost:5000 (Express) → outbound TLS to Aiven MySQL and "
        "Gemini. Production would place the app behind HTTPS reverse proxy. See also frontend/architecture.html.",
    )
    set_runs(
        doc.paragraphs[117],
        "Detailed design artefacts (include screenshots/exports from the project docs pages in the Appendix):\n\n"
        "Entity Relationship: users (students/admins) 1—1 employers; employers 1—N jobs; users N—M jobs via "
        "applications and bookmarks; jobs 1—N job_views; users 1—N notifications. Key statuses: job "
        "(pending/active/closed/rejected); application (pending/reviewed/shortlisted/rejected/hired); employer "
        "verification (pending/approved/rejected); account_status (active/removed).\n\n"
        "Use cases: Student — Register, Search Jobs, Apply, Bookmark, Manage Profile, View Notifications; Employer — "
        "Register Company, Post/Edit Job, Manage Applicants, View Insights; Admin — Moderate Jobs/Users, Verify "
        "Employers, View Analytics. See frontend/usecase.html and workflow.html.\n\n"
        "Implementation notes: Password hashing with bcrypt; JWT in Authorization header; employer-created jobs "
        "default to pending; closeExpiredJobs utility closes past-deadline active jobs; notifications created on "
        "status change and when new active jobs match student interest_fields; AI search/chat quota enforced in "
        "backend/routes/chat.js and utils/aiSearch.js.\n\n"
        "Diagrams recommended to paste here: ERD (erd.html), Architecture (architecture.html), Workflow "
        "(workflow.html), Storyboard (storyboard.html).",
    )
    for i in range(119, 127):
        clear_para(doc.paragraphs[i])

    # Testing
    set_runs(
        doc.paragraphs[129],
        "Test approach: Manual functional testing against role-based requirements, plus exploratory UI testing and "
        "regression checks after major features (notifications, resume upload, admin soft-remove, employer dashboard). "
        "A formal test specification table is recommended in Appendix B using RP Test Specification Template columns "
        "(Test ID, Description, Preconditions, Steps, Expected, Actual, Status).\n\n"
        "Sample test cases (summary):\n"
        "T01 Student register/login — valid credentials issue JWT and redirect; invalid password rejected.\n"
        "T02 Browse/search/filter jobs — only active non-expired jobs appear to students.\n"
        "T03 Apply with cover letter + resume PDF — application stored; duplicate apply blocked.\n"
        "T04 Bookmark add/remove — appears on profile; unique constraint enforced.\n"
        "T05 Employer post job — job created as pending; not visible publicly until admin approve.\n"
        "T06 Admin approve/reject job — status updates; rejection reason stored; alerts fire on approve for matching interests.\n"
        "T07 Employer update applicant status — student receives application_status notification.\n"
        "T08 Admin soft-remove user — login shows “This account is invalid.”\n"
        "T09 Expired deadline — job auto-closed; apply blocked; UI shows job removed/unavailable messaging where applicable.\n"
        "T10 Admin analytics — stats/charts load (users, jobs, applications, views).\n"
        "T11 AI chat (if key set) — authenticated request succeeds within rate limits; unauthenticated rejected.\n\n"
        "Environment: Windows 10, Node.js, MySQL on Aiven, Chrome/Edge. Defects found during development (e.g. admin "
        "page spinner, employer status breakdown zeros) were fixed before final demo.",
    )
    for i in [132, 135, 136, 139, 140, 141, 142]:
        clear_para(doc.paragraphs[i])

    # 6.1 User docs
    set_runs(
        doc.paragraphs[146],
        "User guide (non-technical):\n\n"
        "Getting started: Open the site → Register as Student or Employer → Login.\n\n"
        "Students: Update Profile (skills, bio, interest fields, resume URL or PDF). Go to Jobs to search/filter. "
        "Open a job → Apply (cover letter + resume) or Bookmark. Check Profile for applications/bookmarks. Use the "
        "notification bell for status updates and job alerts. Optional: use AI chat for search help if enabled.\n\n"
        "Employers: Complete company profile and UEN; wait for admin verification if required. Post Job with details "
        "and deadline. On Employer dashboard, review applicants and set status (shortlisted/rejected/hired). Use filters "
        "on My Jobs (Approved, Open, Past, Pending, Rejected). Check insights for views/applications and closing soon.\n\n"
        "Admins: Open Admin dashboard for stats and charts. Approve/reject jobs and employers with reasons. Manage users "
        "(filter by role; remove invalid accounts with a reason). Monitor applications and views analytics.\n\n"
        "Demo accounts (confirm against current seed before submission):\n"
        "Admin: admin@jobportal.com\n"
        "Employer: employer@techcorp.com\n"
        "Students: register a new account for testing.",
    )

    # 6.2 Technical docs (insert body after heading)
    tech_text = (
        "Installation / technical manual:\n\n"
        "Prerequisites: Node.js (LTS), npm, MySQL access (local or Aiven), Git.\n\n"
        "1. Clone repository: git clone https://github.com/wohenleii/jobportal.git\n"
        "2. Install dependencies from the project root (npm install) as documented in README.md.\n"
        "3. Configure backend/.env: DB_HOST, DB_USER, DB_PASSWORD, DB_NAME, DB_PORT, JWT_SECRET; optional "
        "GEMINI_API_KEY, GEMINI_MODEL, rate-limit vars.\n"
        "4. Initialise schema/seed: npm run setup-db (runs backend/scripts/setup-db.js). Apply incremental "
        "migrations in backend/scripts/ if upgrading an existing DB (notifications, user removal, UEN, etc.).\n"
        "5. Start: npm run dev (nodemon) or npm start. Open http://localhost:5000.\n\n"
        "Key folders: backend/server.js entry; backend/routes/* APIs; backend/middleware/auth.js; "
        "frontend/* static UI; database/schema.sql & seed.sql.\n"
        "Security notes: Never commit .env or API keys; serve Gemini only via /api/chat; use strong JWT_SECRET in production."
    )
    insert_after(doc.paragraphs[148], tech_text)

    # Gen AI table
    t = doc.tables[1]
    rows_data = [
        (
            "Cursor (AI coding assistant)",
            "Used to help implement and debug features (e.g. notifications, resume upload, admin/employer dashboard fixes), "
            "draft explanations, and assist in completing this report draft. All generated code/text was reviewed, tested, "
            "and edited by the team before inclusion.",
        ),
        (
            "Google Gemini (in-product feature)",
            "Integrated as an optional in-app chat/smart search assistant via server-side API. Outputs support job search "
            "help for logged-in users and are rate-limited; not used as a substitute for academic writing authenticity.",
        ),
        (
            "[Optional: ChatGPT / other — fill if used]",
            "[Describe how outputs were used, e.g. brainstorming test cases, rephrasing paragraphs. State that final wording was verified by the team.]",
        ),
    ]
    for i, (tool, how) in enumerate(rows_data):
        row = t.rows[i + 1]
        row.cells[0].text = tool
        row.cells[1].text = how

    p = find_para(doc, startswith="Section Explanation: You should have used GitHub")
    if p:
        set_runs(
            p,
            "The JobPortal source code is hosted on GitHub. Collaborators should ensure Gen.SOI-Project@rp.edu.sg "
            "is added as required by the module instructions (GitHub collaborator / GitLab member). Provide a screenshot "
            "of repository access settings in this section.",
        )

    p = find_para(doc, equals="Indicate your Git URL.")
    if p:
        set_runs(p, "Git repository URL for this FYP:")

    p = find_para(doc, startswith="Git URL:")
    if p:
        set_runs(p, "Git URL: https://github.com/wohenleii/jobportal.git")

    p = find_para(doc, equals="GitLab Screenshot")
    if p:
        set_runs(p, "GitHub Screenshot: [PASTE screenshot of repository page / collaborators here]")

    p = find_para(doc, startswith="Section Explanation: This is the section where you should state")
    if p:
        set_runs(
            p,
            "JobPortal successfully demonstrates a complete multi-role job portal spanning requirements, design, "
            "implementation, and testing. The main takeaway is that a moderated marketplace with clear student, "
            "employer, and admin workflows can reduce job-search friction while improving trust and visibility.",
        )

    for label, content in [
        (
            "Summary of Accomplishments",
            "Summary of Accomplishments: Delivered authentication and RBAC; job search/apply/bookmark; employer posting "
            "and applicant management; admin moderation and analytics; notifications and interest-based alerts; resume "
            "upload; employer verification; expired-job handling; optional Gemini AI assist; and project documentation pages.",
        ),
        (
            "Individual reflections",
            "Individual reflections:\n"
            "[Member 1 — FILL IN: what you learned, challenges, contribution]\n"
            "[Member 2 — FILL IN]\n"
            "[Member 3 — FILL IN]\n"
            "[Member 4 — FILL IN]",
        ),
        (
            "Future Work/Recommendations",
            "Future Work/Recommendations: Email/SMS notification channels; password reset and 2FA; public company pages; "
            "richer AI matching of resumes to jobs; automated test suite (Jest/Supertest + Playwright); CI/CD deployment "
            "to a cloud host with HTTPS; accessibility audit; and mobile-responsive polish.",
        ),
    ]:
        p = find_para(doc, equals=label)
        if p:
            set_runs(p, content)

    p = find_para(doc, startswith="Section Explanation: Acknowledge any work")
    if p:
        set_runs(
            p,
            "References (IEEE style):\n"
            "[1] Express.js documentation. [Online]. Available: https://expressjs.com/\n"
            "[2] Oracle, MySQL Reference Manual. [Online]. Available: https://dev.mysql.com/doc/\n"
            "[3] Bootstrap 5 documentation. [Online]. Available: https://getbootstrap.com/docs/5.0/\n"
            "[4] JSON Web Token (JWT) introduction. [Online]. Available: https://jwt.io/introduction\n"
            "[5] Google AI Gemini API documentation. [Online]. Available: https://ai.google.dev/\n"
            "[6] Chart.js documentation. [Online]. Available: https://www.chartjs.org/docs/\n"
            "[7] Aiven for MySQL documentation. [Online]. Available: https://aiven.io/docs/products/mysql\n"
            "[8] M. Fowler, Patterns of Enterprise Application Architecture. Boston, MA: Addison-Wesley, 2002.",
        )

    for p in doc.paragraphs:
        t = p.text.strip()
        if t in ("IEEE citation style", "APA citation style", "MLA citation style") or t.startswith(
            "[1] D. Ingre"
        ) or t.startswith("Usage:") or t.startswith("Dubeck") or t.startswith("James, Nancy") or "Kirk and Spock" in t:
            clear_para(p)

    p = find_para(doc, startswith="Section Explanation: This is for you to include materials")
    if p:
        set_runs(
            p,
            "Suggested appendices:\n"
            "Appendix A — Selected source listings or link to GitHub tree (do not dump entire codebase).\n"
            "Appendix B — Full test specification table (RP template).\n"
            "Appendix C — Screenshots: student apply flow, employer dashboard, admin analytics, notifications.\n"
            "Appendix D — Exported diagrams: ERD, architecture, use cases, Gantt, workflow.\n"
            "Appendix E — Environment variable template (.env.example) without secrets.\n"
            "Appendix F — Meeting minutes / fair work allocation evidence.",
        )
    p = find_para(doc, startswith="Appendices should be headed")
    if p:
        clear_para(p)

    p = find_para(doc, startswith="Section Explanation: Develop a poster")
    if p:
        set_runs(
            p,
            "Embed the project poster PPT/PDF object here. Poster should highlight: JobPortal title, problem & solution, "
            "three-role architecture diagram, key features (search/apply, employer pipeline, admin analytics, notifications/AI), "
            "tech stack (Node/Express, MySQL, Bootstrap), and team names/IDs. File suggestion: JobPortal_Poster.pptx "
            "(create separately and Insert → Object in Word).",
        )

    doc.save(str(OUT1))
    shutil.copy2(OUT1, OUT2)
    print("Saved:", OUT1)
    print("Saved:", OUT2)


if __name__ == "__main__":
    main()
